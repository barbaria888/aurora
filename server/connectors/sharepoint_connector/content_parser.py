"""Parse SharePoint pages and documents into structured text / markdown."""

from __future__ import annotations

import io
import logging
from typing import Any, Dict, List

from bs4 import BeautifulSoup, NavigableString, Tag

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# SharePoint page -> markdown
# ---------------------------------------------------------------------------

def sharepoint_page_to_markdown(canvas_layout: Dict[str, Any]) -> str:
    """Convert a SharePoint page canvas layout to lightweight Markdown.

    Iterates over ``horizontalSections`` -> ``columns`` -> ``webparts`` and
    extracts ``innerHtml`` from text web parts, converting the HTML to markdown.

    Args:
        canvas_layout: The ``canvasLayout`` dict from the Graph API page response.

    Returns:
        Markdown string representation of the page content.
    """
    if not canvas_layout:
        return ""

    parts: List[str] = []
    sections = canvas_layout.get("horizontalSections", [])
    for section in sections:
        for column in section.get("columns", []):
            for webpart in column.get("webparts", []):
                inner_html = (webpart.get("data") or {}).get("innerHtml", "")
                if inner_html:
                    md = _html_to_markdown(inner_html)
                    if md.strip():
                        parts.append(md)

    return "\n\n".join(parts).strip()


def _html_to_markdown(html_content: str) -> str:
    """Convert an HTML fragment to simple Markdown."""
    soup = BeautifulSoup(html_content or "", "html.parser")
    lines: List[str] = []

    def render_element(element: Any, indent: int = 0) -> None:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                lines.append(text)
            return

        if not isinstance(element, Tag):
            return

        tag_name = element.name.lower() if element.name else ""

        if tag_name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(tag_name[1])
            heading = element.get_text(" ", strip=True)
            if heading:
                lines.append(f"{'#' * level} {heading}")
                lines.append("")
            return

        if tag_name == "p":
            text = element.get_text(" ", strip=True)
            if text:
                lines.append(text)
                lines.append("")
            return

        if tag_name == "pre":
            code_text = element.get_text("\n", strip=False).rstrip()
            lines.append("```")
            if code_text:
                lines.append(code_text)
            lines.append("```")
            lines.append("")
            return

        if tag_name in {"ul", "ol"}:
            ordered = tag_name == "ol"
            items = element.find_all("li", recursive=False)
            for idx, li in enumerate(items, start=1):
                prefix = f"{idx}." if ordered else "-"
                item_text = li.get_text(" ", strip=True)
                lines.append(f"{'  ' * indent}{prefix} {item_text}")
                for child_list in li.find_all(["ul", "ol"], recursive=False):
                    render_element(child_list, indent + 1)
            lines.append("")
            return

        if tag_name == "code":
            code_text = element.get_text(" ", strip=True)
            if code_text:
                lines.append(f"`{code_text}`")
            return

        if tag_name == "table":
            _render_table(element, lines)
            return

        for child in element.children:
            render_element(child, indent)

    for child in soup.contents:
        render_element(child)

    return _collapse_blank_lines(lines)


def _render_table(table: Tag, lines: List[str]) -> None:
    """Render an HTML table as a markdown table."""
    rows = table.find_all("tr")
    if not rows:
        return

    parsed_rows: List[List[str]] = []
    for row in rows:
        cells = row.find_all(["th", "td"])
        parsed_rows.append([cell.get_text(" ", strip=True).replace("|", "\\|") for cell in cells])

    if not parsed_rows:
        return

    # First row as header
    header = parsed_rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row_data in parsed_rows[1:]:
        # Pad row to header length if necessary
        padded = row_data + [""] * (len(header) - len(row_data))
        lines.append("| " + " | ".join(padded[: len(header)]) + " |")
    lines.append("")


# ---------------------------------------------------------------------------
# Document text extraction
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file using python-docx.

    Args:
        file_bytes: Raw bytes of the .docx file.

    Returns:
        Extracted text content.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is not installed; cannot extract .docx text")
        return ""

    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_texts:
                    parts.append(" | ".join(cell_texts))
        return "\n\n".join(parts)
    except Exception as exc:
        logger.error("Failed to extract text from .docx: %s", exc)
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Extracted text content.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("pypdf is not installed; cannot extract PDF text")
        return ""

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_text: List[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(text.strip())
        return "\n\n".join(pages_text)
    except Exception as exc:
        logger.error("Failed to extract text from PDF: %s", exc)
        return ""


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract text from an .xlsx file using openpyxl.

    Args:
        file_bytes: Raw bytes of the .xlsx file.

    Returns:
        Extracted text content (one sheet per section).
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.error("openpyxl is not installed; cannot extract .xlsx text")
        return ""

    try:
        with load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True) as wb:
            parts: List[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows_text: List[str] = []
                for row in ws.iter_rows(values_only=True):
                    cell_values = [str(cell) if cell is not None else "" for cell in row]
                    row_str = "\t".join(cell_values).strip()
                    if row_str:
                        rows_text.append(row_str)
                if rows_text:
                    parts.append(f"## {sheet_name}\n\n" + "\n".join(rows_text))
            return "\n\n".join(parts)
    except Exception as exc:
        logger.error("Failed to extract text from .xlsx: %s", exc)
        return ""


def extract_document_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch document text extraction based on file extension.

    Args:
        file_bytes: Raw bytes of the file.
        filename: Original filename (used to detect extension).

    Returns:
        Extracted text content.
    """
    if not filename:
        return ""

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("docx",):
        return extract_text_from_docx(file_bytes)
    elif ext in ("pdf",):
        return extract_text_from_pdf(file_bytes)
    elif ext in ("xlsx",):
        return extract_text_from_xlsx(file_bytes)
    elif ext in ("txt", "md", "csv", "log", "json", "xml", "yaml", "yml"):
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception as exc:
            logger.error("Failed to decode text file %s: %s", filename, exc)
            return ""
    else:
        logger.info(
            "Unsupported file extension '%s' for text extraction from '%s'",
            ext,
            filename,
        )
        return f"[Unsupported file type: .{ext}] Cannot extract text from '{filename}'."


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _collapse_blank_lines(lines: List[str]) -> str:
    """Collapse consecutive blank lines and strip trailing whitespace."""
    cleaned: List[str] = []
    previous_blank = False
    for line in lines:
        if not line.strip():
            if not previous_blank:
                cleaned.append("")
            previous_blank = True
        else:
            cleaned.append(line)
            previous_blank = False
    return "\n".join(cleaned).strip()
